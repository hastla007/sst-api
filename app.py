from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Header, Depends, BackgroundTasks, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import JSONResponse, PlainTextResponse, FileResponse, StreamingResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from typing import Optional, List, Dict, Any, Union
from datetime import datetime, timedelta
from sqlalchemy import select, or_, func as sql_func, text
from sqlalchemy.orm import Session
import whisper
import tempfile
import os
import logging
import hashlib
import json
import uuid
import time
import io
from collections import defaultdict, deque
from functools import wraps
import asyncio
import aiohttp
from urllib.parse import urlparse
import numpy as np

# Database imports
from database import get_db_session, init_db_async, check_db_connection, async_engine
from models import (
    Organization, Project, APIKey, Transcription, UsageLog, WebhookLog,
    SubscriptionTier, TranscriptionStatus
)

# Third-party imports
try:
    import redis
    REDIS_AVAILABLE = True
except ImportError:
    REDIS_AVAILABLE = False
    logging.warning("Redis not available. Caching and advanced features will be disabled.")

try:
    from celery import Celery
    CELERY_AVAILABLE = True
except ImportError:
    CELERY_AVAILABLE = False
    logging.warning("Celery not available. Async queue features will be disabled.")

try:
    from prometheus_client import Counter, Histogram, Gauge, generate_latest, CONTENT_TYPE_LATEST
    PROMETHEUS_AVAILABLE = True
except ImportError:
    PROMETHEUS_AVAILABLE = False
    logging.warning("Prometheus client not available. Metrics will be disabled.")

# Cloud storage imports
try:
    import boto3
    from botocore.exceptions import ClientError
    S3_AVAILABLE = True
except ImportError:
    S3_AVAILABLE = False

try:
    from google.cloud import storage as gcs_storage
    GCS_AVAILABLE = True
except ImportError:
    GCS_AVAILABLE = False

try:
    from azure.storage.blob import BlobServiceClient
    AZURE_AVAILABLE = True
except ImportError:
    AZURE_AVAILABLE = False

# Speaker diarization
try:
    from pyannote.audio import Pipeline
    DIARIZATION_AVAILABLE = True
except ImportError:
    DIARIZATION_AVAILABLE = False
    logging.warning("Pyannote.audio not available. Speaker diarization will be disabled.")

# Enhanced export formats
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Translation and Sentiment
try:
    from transformers import (
        MarianMTModel, MarianTokenizer,
        pipeline as transformers_pipeline
    )
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    logging.warning("Transformers not available. Translation and sentiment analysis will be disabled.")

# Setup logging with custom filter for correlation_id
class CorrelationIdFilter(logging.Filter):
    def filter(self, record):
        if not hasattr(record, 'correlation_id'):
            record.correlation_id = 'N/A'
        return True

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - [%(correlation_id)s] %(message)s'
)
logger = logging.getLogger(__name__)
logger.addFilter(CorrelationIdFilter())

# Configuration
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB limit
SUPPORTED_LANGUAGES = {
    "en", "zh", "de", "es", "ru", "ko", "fr", "ja", "pt", "tr", "pl", "ca", "nl",
    "ar", "sv", "it", "id", "hi", "fi", "vi", "he", "uk", "el", "ms", "cs", "ro",
    "da", "hu", "ta", "no", "th", "ur", "hr", "bg", "lt", "la", "mi", "ml", "cy",
    "sk", "te", "fa", "lv", "bn", "sr", "az", "sl", "kn", "et", "mk", "br", "eu",
    "is", "hy", "ne", "mn", "bs", "kk", "sq", "sw", "gl", "mr", "pa", "si", "km",
    "sn", "yo", "so", "af", "oc", "ka", "be", "tg", "sd", "gu", "am", "yi", "lo",
    "uz", "fo", "ht", "ps", "tk", "nn", "mt", "sa", "lb", "my", "bo", "tl", "mg",
    "as", "tt", "haw", "ln", "ha", "ba", "jw", "su"
}
SUPPORTED_AUDIO_TYPES = {
    "audio/mpeg", "audio/mp3", "audio/wav", "audio/wave", "audio/x-wav",
    "audio/mp4", "audio/m4a", "audio/x-m4a", "audio/ogg", "audio/webm",
    "audio/flac", "audio/aac"
}

# Security Configuration
API_KEYS = set(os.getenv("API_KEYS", "").split(",")) if os.getenv("API_KEYS") else set()
ENABLE_AUTH = len(API_KEYS) > 0
RATE_LIMIT_REQUESTS = int(os.getenv("RATE_LIMIT_REQUESTS", "10"))
RATE_LIMIT_WINDOW = int(os.getenv("RATE_LIMIT_WINDOW", "60"))  # seconds

# Redis Configuration
REDIS_HOST = os.getenv("REDIS_HOST", "redis")
REDIS_PORT = int(os.getenv("REDIS_PORT", "6379"))
REDIS_DB = int(os.getenv("REDIS_DB", "0"))
CACHE_TTL = int(os.getenv("CACHE_TTL", "3600"))  # 1 hour default

# Celery Configuration
CELERY_BROKER = os.getenv("CELERY_BROKER", f"redis://{REDIS_HOST}:{REDIS_PORT}/1")
CELERY_BACKEND = os.getenv("CELERY_BACKEND", f"redis://{REDIS_HOST}:{REDIS_PORT}/2")

# GPU Configuration
USE_GPU = os.getenv("USE_GPU", "false").lower() == "true"
DEVICE = "cuda" if USE_GPU else "cpu"

# Cloud Storage Configuration
AWS_ACCESS_KEY = os.getenv("AWS_ACCESS_KEY_ID", "")
AWS_SECRET_KEY = os.getenv("AWS_SECRET_ACCESS_KEY", "")
AWS_BUCKET = os.getenv("AWS_BUCKET", "stt-audio-files")
AWS_REGION = os.getenv("AWS_REGION", "us-east-1")

GCS_BUCKET = os.getenv("GCS_BUCKET", "stt-audio-files")
GCS_CREDENTIALS_FILE = os.getenv("GCS_CREDENTIALS_FILE", "")

AZURE_CONNECTION_STRING = os.getenv("AZURE_STORAGE_CONNECTION_STRING", "")
AZURE_CONTAINER = os.getenv("AZURE_CONTAINER", "stt-audio-files")

# Diarization Configuration
DIARIZATION_MODEL = os.getenv("DIARIZATION_MODEL", "pyannote/speaker-diarization")
HF_TOKEN = os.getenv("HUGGINGFACE_TOKEN", "")  # Required for pyannote models

# Initialize Redis client
redis_client = None
if REDIS_AVAILABLE:
    try:
        redis_client = redis.Redis(
            host=REDIS_HOST,
            port=REDIS_PORT,
            db=REDIS_DB,
            decode_responses=True,
            socket_connect_timeout=5
        )
        redis_client.ping()
        logger.info("Redis connection established")
    except Exception as e:
        logger.warning(f"Failed to connect to Redis: {e}")
        redis_client = None

# Initialize Celery
celery_app = None
if CELERY_AVAILABLE and redis_client:
    celery_app = Celery(
        "stt_tasks",
        broker=CELERY_BROKER,
        backend=CELERY_BACKEND
    )
    celery_app.conf.update(
        task_serializer='json',
        accept_content=['json'],
        result_serializer='json',
        timezone='UTC',
        enable_utc=True,
        task_acks_late=True,
        task_reject_on_worker_lost=True,
    )

# Prometheus Metrics
if PROMETHEUS_AVAILABLE:
    request_counter = Counter('stt_requests_total', 'Total requests', ['endpoint', 'status'])
    request_duration = Histogram('stt_request_duration_seconds', 'Request duration', ['endpoint'])
    transcription_duration = Histogram('stt_transcription_duration_seconds', 'Transcription duration')
    active_requests = Gauge('stt_active_requests', 'Number of active requests')
    cache_hits = Counter('stt_cache_hits_total', 'Total cache hits')
    cache_misses = Counter('stt_cache_misses_total', 'Total cache misses')
    diarization_counter = Counter('stt_diarization_total', 'Total diarization requests')
    translation_counter = Counter('stt_translation_total', 'Total translation requests', ['target_lang'])

# In-memory rate limiting (fallback when Redis is not available)
rate_limit_storage = defaultdict(lambda: deque())

# Usage analytics storage
usage_stats = {
    "total_requests": 0,
    "successful_requests": 0,
    "failed_requests": 0,
    "total_duration": 0.0,
    "cache_hits": 0,
    "cache_misses": 0,
}

app = FastAPI(
    title="Enterprise STT Service API v3.0",
    version="3.0.0",
    description="""
    ## Enterprise-Grade Speech-to-Text API with Advanced Features

    ### Core Features:
    - 🔐 Multi-Tenant Organization Management
    - 🎯 Speaker Diarization (Who spoke when)
    - 🌍 Auto Language Detection + 99 Language Support
    - 📊 Confidence Scores per Segment
    - 💬 Custom Vocabulary & Prompts
    - 🔄 Translation to Multiple Languages
    - 🧠 Sentiment Analysis
    - ⚡ Real-Time Streaming Transcription
    - 📦 Batch Processing
    - ☁️ Cloud Storage Integration (S3, GCS, Azure)
    - 📝 Enhanced Export Formats (JSON, SRT, VTT, DOCX, PDF, TXT)
    - 💾 Persistent Database Storage
    - 🔍 Full-Text Search
    - 💰 Usage Quotas & Billing
    - 🔔 Webhook Support with Retry Logic
    - 🚀 Async Queue Processing
    - 💾 Redis Caching
    - 🎯 GPU Support
    - 📊 Prometheus Metrics
    - 🔍 Request Tracking

    ### Authentication:
    Pass your API key in the `X-API-Key` header.

    ### Rate Limits:
    - Default: 10 requests per 60 seconds
    - Configurable per organization
    """,
    docs_url="/docs",
    redoc_url="/redoc",
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "*").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

# Load Whisper model
MODEL_SIZE = os.getenv("WHISPER_MODEL", "base")
logger.info(f"Loading Whisper model: {MODEL_SIZE} on device: {DEVICE}")
model = whisper.load_model(MODEL_SIZE, device=DEVICE)
logger.info("Model loaded successfully")

# Load diarization pipeline if available
diarization_pipeline = None
if DIARIZATION_AVAILABLE and HF_TOKEN:
    try:
        diarization_pipeline = Pipeline.from_pretrained(
            DIARIZATION_MODEL,
            use_auth_token=HF_TOKEN
        )
        if USE_GPU:
            import torch
            diarization_pipeline.to(torch.device("cuda"))
        logger.info("Diarization pipeline loaded successfully")
    except Exception as e:
        logger.warning(f"Failed to load diarization pipeline: {e}")
        diarization_pipeline = None

# Load sentiment analysis pipeline
sentiment_analyzer = None
if TRANSFORMERS_AVAILABLE:
    try:
        sentiment_analyzer = transformers_pipeline("sentiment-analysis", model="distilbert-base-uncased-finetuned-sst-2-english")
        logger.info("Sentiment analyzer loaded successfully")
    except Exception as e:
        logger.warning(f"Failed to load sentiment analyzer: {e}")

# Translation models cache
translation_models = {}


@app.on_event("startup")
async def startup_event():
    """Initialize database and services on startup"""
    try:
        await init_db_async()
        logger.info("Database initialized")
    except Exception as e:
        logger.error(f"Failed to initialize database: {e}")


# Middleware for correlation ID
@app.middleware("http")
async def add_correlation_id(request: Request, call_next):
    correlation_id = request.headers.get("X-Correlation-ID", str(uuid.uuid4()))
    request.state.correlation_id = correlation_id

    # Add correlation ID to logger context
    old_factory = logging.getLogRecordFactory()

    def record_factory(*args, **kwargs):
        record = old_factory(*args, **kwargs)
        record.correlation_id = getattr(record, 'correlation_id', correlation_id)
        return record

    logging.setLogRecordFactory(record_factory)

    response = await call_next(request)
    response.headers["X-Correlation-ID"] = correlation_id

    logging.setLogRecordFactory(old_factory)
    return response


# Authentication dependency - now checks database
async def verify_api_key(
    x_api_key: Optional[str] = Header(None),
    db: Session = Depends(get_db_session)
) -> Optional[APIKey]:
    """Verify API key and return associated key object"""
    if not ENABLE_AUTH:
        return None

    if not x_api_key:
        raise HTTPException(
            status_code=401,
            detail="Missing API key. Please provide X-API-Key header."
        )

    # Check database for API key
    try:
        result = await db.execute(
            select(APIKey).where(APIKey.key == x_api_key, APIKey.is_active == True)
        )
        api_key_obj = result.scalar_one_or_none()

        if not api_key_obj:
            raise HTTPException(
                status_code=403,
                detail="Invalid or inactive API key"
            )

        # Update last used
        api_key_obj.last_used_at = datetime.utcnow()
        api_key_obj.total_requests += 1
        await db.commit()

        return api_key_obj

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error verifying API key: {e}")
        raise HTTPException(status_code=500, detail="Authentication error")


# Rate limiting dependency with database integration
async def rate_limiter(
    request: Request,
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    identifier = api_key.id if api_key else request.client.host
    current_time = time.time()

    # Get rate limit (from API key or default)
    rate_limit = api_key.rate_limit if api_key else RATE_LIMIT_REQUESTS

    if redis_client:
        # Use Redis for distributed rate limiting
        key = f"rate_limit:{identifier}"
        try:
            pipe = redis_client.pipeline()
            pipe.zadd(key, {str(current_time): current_time})
            pipe.zremrangebyscore(key, 0, current_time - RATE_LIMIT_WINDOW)
            pipe.zcard(key)
            pipe.expire(key, RATE_LIMIT_WINDOW)
            results = pipe.execute()
            request_count = results[2]
        except Exception as e:
            logger.error(f"Redis rate limiting error: {e}")
            request_count = 0
    else:
        # Fallback to in-memory rate limiting
        requests = rate_limit_storage[identifier]
        # Remove old requests
        while requests and requests[0] < current_time - RATE_LIMIT_WINDOW:
            requests.popleft()
        requests.append(current_time)
        request_count = len(requests)

    if request_count > rate_limit:
        raise HTTPException(
            status_code=429,
            detail=f"Rate limit exceeded. Maximum {rate_limit} requests per {RATE_LIMIT_WINDOW} seconds."
        )

    return True


# Utility functions
def calculate_file_hash(content: bytes) -> str:
    """Calculate SHA256 hash of file content"""
    return hashlib.sha256(content).hexdigest()


def validate_webhook_url(url: str) -> bool:
    """Validate webhook URL format"""
    try:
        result = urlparse(url)
        return all([result.scheme in ['http', 'https'], result.netloc])
    except Exception:
        return False


def get_cached_result(cache_key: str) -> Optional[Dict]:
    """Get cached transcription result"""
    if not redis_client:
        return None

    try:
        cached = redis_client.get(f"transcription:{cache_key}")
        if cached:
            if PROMETHEUS_AVAILABLE:
                cache_hits.inc()
            usage_stats["cache_hits"] += 1
            return json.loads(cached)
    except Exception as e:
        logger.error(f"Cache retrieval error: {e}")

    if PROMETHEUS_AVAILABLE:
        cache_misses.inc()
    usage_stats["cache_misses"] += 1
    return None


def set_cached_result(cache_key: str, result: Dict):
    """Cache transcription result"""
    if not redis_client:
        return

    try:
        redis_client.setex(
            f"transcription:{cache_key}",
            CACHE_TTL,
            json.dumps(result)
        )
    except Exception as e:
        logger.error(f"Cache storage error: {e}")


def format_timestamp(seconds: float) -> str:
    """Format seconds to SRT/VTT timestamp"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def segments_to_srt(segments: List[Dict]) -> str:
    """Convert segments to SRT format"""
    srt_content = []
    for i, segment in enumerate(segments, 1):
        start = format_timestamp(segment['start'])
        end = format_timestamp(segment['end'])
        text = segment['text'].strip()
        speaker = f"[{segment.get('speaker', 'Unknown')}] " if segment.get('speaker') else ""
        srt_content.append(f"{i}\n{start} --> {end}\n{speaker}{text}\n")
    return "\n".join(srt_content)


def segments_to_vtt(segments: List[Dict]) -> str:
    """Convert segments to WebVTT format"""
    vtt_content = ["WEBVTT\n"]
    for segment in segments:
        start = format_timestamp(segment['start']).replace(',', '.')
        end = format_timestamp(segment['end']).replace(',', '.')
        text = segment['text'].strip()
        speaker = f"[{segment.get('speaker', 'Unknown')}] " if segment.get('speaker') else ""
        vtt_content.append(f"{start} --> {end}\n{speaker}{text}\n")
    return "\n".join(vtt_content)


def segments_to_txt(segments: List[Dict]) -> str:
    """Convert segments to plain text"""
    lines = []
    for segment in segments:
        speaker = f"[{segment.get('speaker', 'Unknown')}] " if segment.get('speaker') else ""
        lines.append(f"{speaker}{segment['text'].strip()}")
    return "\n".join(lines)


def segments_to_docx(segments: List[Dict], filename: str = "transcription.docx") -> str:
    """Convert segments to DOCX format"""
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=503, detail="DOCX export not available")

    doc = Document()
    doc.add_heading('Transcription', 0)

    for segment in segments:
        timestamp = f"[{format_timestamp(segment['start'])} --> {format_timestamp(segment['end'])}]"
        speaker = f"Speaker: {segment.get('speaker', 'Unknown')}" if segment.get('speaker') else ""

        p = doc.add_paragraph()
        p.add_run(timestamp).bold = True
        if speaker:
            p.add_run(f"\n{speaker}\n").italic = True
        p.add_run(segment['text'].strip())

    temp_path = f"/tmp/{filename}"
    doc.save(temp_path)
    return temp_path


def segments_to_pdf(segments: List[Dict], filename: str = "transcription.pdf") -> str:
    """Convert segments to PDF format"""
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=503, detail="PDF export not available")

    temp_path = f"/tmp/{filename}"
    doc = SimpleDocTemplate(temp_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Title
    title = Paragraph("Transcription", styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))

    # Content
    for segment in segments:
        timestamp = f"[{format_timestamp(segment['start'])} --> {format_timestamp(segment['end'])}]"
        speaker = f"Speaker: {segment.get('speaker', 'Unknown')}" if segment.get('speaker') else ""

        text_content = f"{timestamp}<br/>{speaker}<br/>{segment['text'].strip()}"
        p = Paragraph(text_content, styles['Normal'])
        story.append(p)
        story.append(Spacer(1, 12))

    doc.build(story)
    return temp_path


async def download_from_url(url: str) -> tuple[bytes, str]:
    """Download file from cloud storage URL or HTTP URL"""
    # S3 URL
    if url.startswith("s3://") and S3_AVAILABLE:
        bucket_name = url.split("/")[2]
        key = "/".join(url.split("/")[3:])
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_ACCESS_KEY,
            aws_secret_access_key=AWS_SECRET_KEY,
            region_name=AWS_REGION
        )
        obj = s3_client.get_object(Bucket=bucket_name, Key=key)
        content = obj['Body'].read()
        filename = key.split("/")[-1]
        return content, filename

    # GCS URL
    elif url.startswith("gs://") and GCS_AVAILABLE:
        bucket_name = url.split("/")[2]
        blob_name = "/".join(url.split("/")[3:])
        client = gcs_storage.Client()
        bucket = client.bucket(bucket_name)
        blob = bucket.blob(blob_name)
        content = blob.download_as_bytes()
        filename = blob_name.split("/")[-1]
        return content, filename

    # Azure URL
    elif "blob.core.windows.net" in url and AZURE_AVAILABLE:
        blob_client = BlobServiceClient.from_connection_string(AZURE_CONNECTION_STRING)
        # Parse URL to get container and blob
        parts = url.split("/")
        container_name = parts[-2]
        blob_name = parts[-1]
        blob = blob_client.get_blob_client(container=container_name, blob=blob_name)
        content = blob.download_blob().readall()
        return content, blob_name

    # HTTP/HTTPS URL
    elif url.startswith(("http://", "https://")):
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise HTTPException(status_code=400, detail=f"Failed to download file from URL: {response.status}")
                content = await response.read()
                filename = url.split("/")[-1].split("?")[0]
                return content, filename
    else:
        raise HTTPException(status_code=400, detail="Unsupported URL format")


async def send_webhook_with_retry(webhook_url: str, data: Dict, transcription_id: str, db: Session):
    """Send webhook with retry logic"""
    from sqlalchemy import select

    # Get or create webhook log
    result = await db.execute(
        select(WebhookLog).where(WebhookLog.transcription_id == transcription_id)
    )
    webhook_log = result.scalar_one_or_none()

    if not webhook_log:
        webhook_log = WebhookLog(
            id=str(uuid.uuid4()),
            transcription_id=transcription_id,
            webhook_url=webhook_url,
            created_at=datetime.utcnow()
        )
        db.add(webhook_log)
        await db.commit()

    max_retries = webhook_log.max_attempts
    retry_delays = [2, 4, 8]  # Exponential backoff: 2s, 4s, 8s

    for attempt in range(max_retries):
        try:
            webhook_log.attempts = attempt + 1
            webhook_log.last_attempt_at = datetime.utcnow()

            async with aiohttp.ClientSession() as session:
                async with session.post(
                    webhook_url,
                    json=data,
                    timeout=aiohttp.ClientTimeout(total=30)
                ) as response:
                    webhook_log.status_code = response.status
                    webhook_log.response_body = await response.text()

                    if response.status < 400:
                        webhook_log.success = True
                        webhook_log.completed_at = datetime.utcnow()
                        await db.commit()
                        logger.info(f"Webhook sent successfully to {webhook_url}")
                        return
                    else:
                        webhook_log.error_message = f"HTTP {response.status}"

        except Exception as e:
            webhook_log.error_message = str(e)
            logger.error(f"Webhook attempt {attempt + 1} failed: {e}")

        # Save attempt
        await db.commit()

        # Wait before retry (except on last attempt)
        if attempt < max_retries - 1:
            await asyncio.sleep(retry_delays[attempt])
            webhook_log.next_retry_at = datetime.utcnow() + timedelta(seconds=retry_delays[min(attempt + 1, len(retry_delays) - 1)])
            await db.commit()

    logger.error(f"Webhook failed after {max_retries} attempts to {webhook_url}")


def perform_diarization(audio_path: str) -> Optional[Dict]:
    """Perform speaker diarization"""
    if not diarization_pipeline:
        return None

    try:
        if PROMETHEUS_AVAILABLE:
            diarization_counter.inc()

        diarization = diarization_pipeline(audio_path)

        speakers_data = []
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            speakers_data.append({
                "start": turn.start,
                "end": turn.end,
                "speaker": speaker
            })

        return {"speakers": speakers_data}

    except Exception as e:
        logger.error(f"Diarization error: {e}")
        return None


def merge_diarization_with_segments(segments: List[Dict], diarization_data: Optional[Dict]) -> List[Dict]:
    """Merge speaker diarization with transcription segments"""
    if not diarization_data:
        return segments

    speakers = diarization_data.get("speakers", [])
    if not speakers:
        return segments

    # Assign speakers to segments based on overlap
    for segment in segments:
        segment_start = segment['start']
        segment_end = segment['end']
        segment_mid = (segment_start + segment_end) / 2

        # Find speaker with maximum overlap
        best_speaker = None
        max_overlap = 0

        for speaker_turn in speakers:
            overlap_start = max(segment_start, speaker_turn['start'])
            overlap_end = min(segment_end, speaker_turn['end'])
            overlap = max(0, overlap_end - overlap_start)

            if overlap > max_overlap:
                max_overlap = overlap
                best_speaker = speaker_turn['speaker']

        if best_speaker:
            segment['speaker'] = best_speaker

    return segments


def translate_text(text: str, target_language: str) -> Optional[str]:
    """Translate text to target language"""
    if not TRANSFORMERS_AVAILABLE:
        return None

    try:
        # Load model for language pair (cache it)
        model_name = f"Helsinki-NLP/opus-mt-en-{target_language}"

        if model_name not in translation_models:
            tokenizer = MarianTokenizer.from_pretrained(model_name)
            model = MarianMTModel.from_pretrained(model_name)
            translation_models[model_name] = (tokenizer, model)

        tokenizer, model = translation_models[model_name]

        # Translate
        inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True, max_length=512)
        outputs = model.generate(**inputs)
        translated = tokenizer.decode(outputs[0], skip_special_tokens=True)

        if PROMETHEUS_AVAILABLE:
            translation_counter.labels(target_lang=target_language).inc()

        return translated

    except Exception as e:
        logger.error(f"Translation error: {e}")
        return None


def analyze_sentiment(text: str) -> Optional[Dict]:
    """Analyze sentiment of text"""
    if not sentiment_analyzer:
        return None

    try:
        # Split into chunks if text is too long
        max_length = 512
        chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]

        results = []
        for chunk in chunks:
            if chunk.strip():
                result = sentiment_analyzer(chunk)[0]
                results.append(result)

        if not results:
            return None

        # Aggregate results
        positive_count = sum(1 for r in results if r['label'] == 'POSITIVE')
        overall_label = 'POSITIVE' if positive_count > len(results) / 2 else 'NEGATIVE'
        overall_score = sum(r['score'] for r in results) / len(results)

        return {
            "label": overall_label,
            "score": overall_score,
            "chunks": results
        }

    except Exception as e:
        logger.error(f"Sentiment analysis error: {e}")
        return None


def transcribe_file(
    file_path: str,
    language: Optional[str] = None,
    initial_prompt: Optional[str] = None,
    return_confidence: bool = False
) -> Dict:
    """Core transcription function with advanced features"""
    transcribe_options = {
        "verbose": True if return_confidence else False,
    }

    if language:
        transcribe_options["language"] = language

    if initial_prompt:
        transcribe_options["initial_prompt"] = initial_prompt

    start_time = time.time()
    result = model.transcribe(file_path, **transcribe_options)
    duration = time.time() - start_time

    if PROMETHEUS_AVAILABLE:
        transcription_duration.observe(duration)

    # Extract confidence scores if requested
    confidence_scores = None
    if return_confidence and "segments" in result:
        confidence_scores = []
        for segment in result["segments"]:
            if "tokens" in segment and segment["tokens"]:
                # Calculate average confidence for segment
                token_probs = []
                for token in segment.get("tokens", []):
                    if isinstance(token, dict) and "probability" in token:
                        token_probs.append(token["probability"])

                avg_confidence = sum(token_probs) / len(token_probs) if token_probs else None
                confidence_scores.append({
                    "segment_id": segment.get("id"),
                    "start": segment.get("start"),
                    "end": segment.get("end"),
                    "confidence": avg_confidence
                })

    return {
        "text": result["text"],
        "language": result.get("language", language),
        "segments": result["segments"],
        "duration": duration,
        "confidence_scores": confidence_scores
    }


# Celery task for async processing with all new features
if celery_app:
    @celery_app.task(name="transcribe_async", bind=True, max_retries=3)
    def transcribe_async_task(
        self,
        file_content: bytes,
        filename: str,
        transcription_id: str,
        language: Optional[str],
        initial_prompt: Optional[str],
        enable_diarization: bool,
        enable_translation: bool,
        target_language: Optional[str],
        enable_sentiment: bool,
        return_confidence: bool,
        webhook_url: Optional[str],
        correlation_id: str
    ):
        """Async transcription task with all features"""
        temp_file_path = None
        try:
            # Save file temporarily
            suffix = os.path.splitext(filename)[1] or ".wav"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            # Transcribe
            result = transcribe_file(temp_file_path, language, initial_prompt, return_confidence)

            # Diarization
            diarization_data = None
            if enable_diarization:
                diarization_data = perform_diarization(temp_file_path)
                if diarization_data:
                    result["segments"] = merge_diarization_with_segments(result["segments"], diarization_data)
                    result["diarization"] = diarization_data

            # Translation
            if enable_translation and target_language:
                translated_text = translate_text(result["text"], target_language)
                if translated_text:
                    result["translation"] = {
                        "target_language": target_language,
                        "text": translated_text
                    }

            # Sentiment analysis
            if enable_sentiment:
                sentiment = analyze_sentiment(result["text"])
                if sentiment:
                    result["sentiment"] = sentiment

            # Clean up
            os.unlink(temp_file_path)

            # Send webhook with retry if provided
            if webhook_url:
                webhook_data = {
                    "correlation_id": correlation_id,
                    "transcription_id": transcription_id,
                    "status": "completed",
                    "result": result
                }

                # Use requests for sync webhook in Celery
                import requests
                for attempt in range(3):
                    try:
                        response = requests.post(webhook_url, json=webhook_data, timeout=30)
                        if response.status_code < 400:
                            break
                    except Exception as e:
                        if attempt == 2:
                            logger.error(f"Webhook failed after 3 attempts: {e}")
                        time.sleep(2 ** attempt)

            return result

        except Exception as e:
            logger.error(f"Async transcription error: {e}")
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)

            # Retry with exponential backoff
            raise self.retry(exc=e, countdown=2 ** self.request.retries)


@app.post(
    "/transcribe",
    dependencies=[Depends(rate_limiter)],
    response_model=None
)
async def transcribe_audio(
    request: Request,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db_session),
    file: Optional[UploadFile] = File(None, description="Audio file (mp3, wav, m4a, etc.) - Max 100MB"),
    file_url: Optional[str] = Form(None, description="URL to audio file (HTTP, S3, GCS, Azure)"),
    language: Optional[str] = Form(None, description="Optional language code (e.g., 'en', 'es', 'de'). Leave empty for auto-detection."),
    initial_prompt: Optional[str] = Form(None, description="Custom vocabulary/context prompt for better accuracy"),
    export_format: Optional[str] = Form("json", description="Export format: json, srt, vtt, txt, docx, pdf"),
    enable_diarization: bool = Form(False, description="Enable speaker diarization (who spoke when)"),
    enable_translation: bool = Form(False, description="Enable translation"),
    target_language: Optional[str] = Form(None, description="Target language for translation (e.g., 'es', 'fr', 'de')"),
    enable_sentiment: bool = Form(False, description="Enable sentiment analysis"),
    return_confidence: bool = Form(True, description="Return confidence scores per segment"),
    webhook_url: Optional[str] = Form(None, description="Webhook URL for async processing"),
    use_cache: bool = Form(True, description="Use cached results if available"),
    project_id: Optional[str] = Form(None, description="Project ID for organization"),
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    """
    **Advanced Transcription Endpoint with All Features**

    Transcribe audio file to text with enterprise features including:
    - Auto language detection or manual language selection
    - Speaker diarization (identify who spoke when)
    - Custom vocabulary/prompts for domain-specific accuracy
    - Translation to multiple languages
    - Sentiment analysis
    - Confidence scores per segment
    - Multiple export formats (JSON, SRT, VTT, TXT, DOCX, PDF)
    - Cloud storage integration (S3, GCS, Azure)
    """
    correlation_id = request.state.correlation_id
    temp_file_path = None
    start_time = time.time()

    try:
        if PROMETHEUS_AVAILABLE:
            active_requests.inc()

        usage_stats["total_requests"] += 1

        # Validate input
        if not file and not file_url:
            raise HTTPException(status_code=400, detail="Either file or file_url must be provided")

        if file and file_url:
            raise HTTPException(status_code=400, detail="Provide either file or file_url, not both")

        # Get file content
        if file_url:
            content, filename = await download_from_url(file_url)
        else:
            filename = file.filename
            content = await file.read()

        file_size = len(content)

        # Validate file size
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE / (1024*1024):.0f}MB"
            )

        if file_size == 0:
            raise HTTPException(status_code=400, detail="Empty file")

        # Validate language code
        if language and language not in SUPPORTED_LANGUAGES:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported language code: {language}. Use /languages endpoint to see supported languages."
            )

        # Validate export format
        if export_format not in ["json", "srt", "vtt", "txt", "docx", "pdf"]:
            raise HTTPException(
                status_code=400,
                detail="Invalid export_format. Must be one of: json, srt, vtt, txt, docx, pdf"
            )

        # Validate webhook URL
        if webhook_url and not validate_webhook_url(webhook_url):
            raise HTTPException(
                status_code=400,
                detail="Invalid webhook URL. Must be a valid HTTP or HTTPS URL."
            )

        # Check quota if API key is provided
        if api_key:
            # Get organization and check quota
            org_result = await db.execute(
                select(Organization).where(Organization.id == api_key.organization_id)
            )
            org = org_result.scalar_one_or_none()

            if org and org.used_quota >= org.monthly_quota:
                raise HTTPException(
                    status_code=429,
                    detail=f"Monthly quota exceeded. Used: {org.used_quota:.2f}/{org.monthly_quota} minutes"
                )

        logger.info(f"Processing file: {filename}, size: {file_size / (1024*1024):.2f}MB")

        # Calculate file hash for caching
        file_hash = calculate_file_hash(content)
        cache_key = f"{file_hash}:{language or 'auto'}:{initial_prompt or ''}"

        # Check cache
        cached_result = None
        if use_cache:
            cached_result = get_cached_result(cache_key)

        if cached_result:
            logger.info("Returning cached result")
            cached_result["correlation_id"] = correlation_id
            cached_result["cached"] = True

            if export_format in ["srt", "vtt", "txt"]:
                if export_format == "srt":
                    content_str = segments_to_srt(cached_result["segments"])
                    media_type = "text/plain"
                elif export_format == "vtt":
                    content_str = segments_to_vtt(cached_result["segments"])
                    media_type = "text/vtt"
                else:
                    content_str = segments_to_txt(cached_result["segments"])
                    media_type = "text/plain"

                return PlainTextResponse(content=content_str, media_type=media_type)

            elif export_format in ["docx", "pdf"]:
                if export_format == "docx":
                    file_path = segments_to_docx(cached_result["segments"], f"{filename}.docx")
                else:
                    file_path = segments_to_pdf(cached_result["segments"], f"{filename}.pdf")

                return FileResponse(
                    path=file_path,
                    filename=os.path.basename(file_path),
                    media_type="application/octet-stream"
                )

            return JSONResponse(content=cached_result)

        # Create transcription record in database
        transcription_id = str(uuid.uuid4())
        transcription = Transcription(
            id=transcription_id,
            project_id=project_id,
            correlation_id=correlation_id,
            filename=filename,
            file_size=file_size,
            file_hash=file_hash,
            status=TranscriptionStatus.PENDING,
            language=language,
            initial_prompt=initial_prompt,
            webhook_url=webhook_url,
            model_used=MODEL_SIZE,
            created_at=datetime.utcnow()
        )
        db.add(transcription)
        await db.commit()

        # Async processing with webhook
        if webhook_url and celery_app:
            transcription.status = TranscriptionStatus.PROCESSING
            transcription.started_at = datetime.utcnow()
            await db.commit()

            task = transcribe_async_task.delay(
                content,
                filename,
                transcription_id,
                language,
                initial_prompt,
                enable_diarization,
                enable_translation,
                target_language,
                enable_sentiment,
                return_confidence,
                webhook_url,
                correlation_id
            )
            logger.info(f"Queued async task: {task.id}")

            return JSONResponse(content={
                "success": True,
                "task_id": task.id,
                "transcription_id": transcription_id,
                "correlation_id": correlation_id,
                "status": "processing",
                "message": "Transcription queued. Results will be sent to webhook URL."
            })

        # Synchronous processing
        transcription.status = TranscriptionStatus.PROCESSING
        transcription.started_at = datetime.utcnow()
        await db.commit()

        suffix = os.path.splitext(filename)[1] or ".wav"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        logger.info(f"Transcribing file: {temp_file_path}")

        # Transcribe with confidence scores
        result = transcribe_file(temp_file_path, language, initial_prompt, return_confidence)

        # Get audio duration
        audio_duration = result["segments"][-1]["end"] if result["segments"] else 0

        # Speaker diarization
        diarization_data = None
        if enable_diarization:
            diarization_data = perform_diarization(temp_file_path)
            if diarization_data:
                result["segments"] = merge_diarization_with_segments(result["segments"], diarization_data)
                result["diarization"] = diarization_data

        # Translation
        if enable_translation and target_language:
            translated_text = translate_text(result["text"], target_language)
            if translated_text:
                result["translation"] = {
                    "target_language": target_language,
                    "text": translated_text
                }

        # Sentiment analysis
        if enable_sentiment:
            sentiment = analyze_sentiment(result["text"])
            if sentiment:
                result["sentiment"] = sentiment

        # Clean up temp file
        os.unlink(temp_file_path)
        temp_file_path = None

        # Update transcription record
        transcription.status = TranscriptionStatus.COMPLETED
        transcription.completed_at = datetime.utcnow()
        transcription.text = result["text"]
        transcription.detected_language = result["language"]
        transcription.segments = result["segments"]
        transcription.confidence_scores = result.get("confidence_scores")
        transcription.speakers = diarization_data
        transcription.translation = result.get("translation")
        transcription.sentiment = result.get("sentiment")
        transcription.processing_time = result["duration"]
        transcription.audio_duration = audio_duration
        await db.commit()

        # Update quota usage
        if api_key:
            org_result = await db.execute(
                select(Organization).where(Organization.id == api_key.organization_id)
            )
            org = org_result.scalar_one_or_none()
            if org:
                org.used_quota += audio_duration / 60  # Convert to minutes
                await db.commit()

            # Log usage
            usage_log = UsageLog(
                id=str(uuid.uuid4()),
                organization_id=api_key.organization_id,
                transcription_id=transcription_id,
                audio_duration=audio_duration / 60,
                processing_time=result["duration"],
                model_used=MODEL_SIZE,
                used_diarization=enable_diarization,
                used_translation=enable_translation,
                used_sentiment=enable_sentiment,
                created_at=datetime.utcnow()
            )
            db.add(usage_log)
            await db.commit()

        # Cache result
        if use_cache:
            set_cached_result(cache_key, result)

        logger.info("Transcription completed successfully")

        # Update stats
        usage_stats["successful_requests"] += 1
        request_duration_time = time.time() - start_time
        usage_stats["total_duration"] += request_duration_time

        if PROMETHEUS_AVAILABLE:
            request_counter.labels(endpoint='transcribe', status='success').inc()
            request_duration.labels(endpoint='transcribe').observe(request_duration_time)

        # Send webhook in background with retry
        if webhook_url:
            webhook_data = {
                "correlation_id": correlation_id,
                "transcription_id": transcription_id,
                "status": "completed",
                "result": result
            }
            background_tasks.add_task(send_webhook_with_retry, webhook_url, webhook_data, transcription_id, db)

        # Prepare response
        response_data = {
            "success": True,
            "transcription_id": transcription_id,
            "text": result["text"],
            "language": result["language"],
            "segments": result["segments"],
            "correlation_id": correlation_id,
            "cached": False,
            "processing_time": result["duration"],
            "audio_duration": audio_duration
        }

        if result.get("confidence_scores"):
            response_data["confidence_scores"] = result["confidence_scores"]

        if diarization_data:
            response_data["diarization"] = diarization_data

        if result.get("translation"):
            response_data["translation"] = result["translation"]

        if result.get("sentiment"):
            response_data["sentiment"] = result["sentiment"]

        # Export in requested format
        if export_format == "srt":
            return PlainTextResponse(
                content=segments_to_srt(result["segments"]),
                media_type="text/plain",
                headers={"X-Correlation-ID": correlation_id}
            )
        elif export_format == "vtt":
            return PlainTextResponse(
                content=segments_to_vtt(result["segments"]),
                media_type="text/vtt",
                headers={"X-Correlation-ID": correlation_id}
            )
        elif export_format == "txt":
            return PlainTextResponse(
                content=segments_to_txt(result["segments"]),
                media_type="text/plain",
                headers={"X-Correlation-ID": correlation_id}
            )
        elif export_format == "docx":
            file_path = segments_to_docx(result["segments"], f"{filename}.docx")
            return FileResponse(
                path=file_path,
                filename=os.path.basename(file_path),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"X-Correlation-ID": correlation_id}
            )
        elif export_format == "pdf":
            file_path = segments_to_pdf(result["segments"], f"{filename}.pdf")
            return FileResponse(
                path=file_path,
                filename=os.path.basename(file_path),
                media_type="application/pdf",
                headers={"X-Correlation-ID": correlation_id}
            )

        return JSONResponse(content=response_data)

    except HTTPException:
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as cleanup_error:
                logger.error(f"Error cleaning up temp file: {cleanup_error}")
        raise

    except Exception as e:
        if PROMETHEUS_AVAILABLE:
            request_counter.labels(endpoint='transcribe', status='error').inc()

        usage_stats["failed_requests"] += 1
        logger.error(f"Transcription error: {str(e)}", exc_info=True)

        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
            except Exception as cleanup_error:
                logger.error(f"Error cleaning up temp file: {cleanup_error}")

        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")

    finally:
        if PROMETHEUS_AVAILABLE:
            active_requests.dec()


@app.websocket("/ws/transcribe")
async def websocket_transcribe(
    websocket: WebSocket,
    api_key: Optional[str] = None
):
    """
    **Real-Time Streaming Transcription via WebSocket**

    Connect to this endpoint to stream audio chunks and receive real-time transcription.

    Protocol:
    1. Send audio chunks as binary data
    2. Send {"action": "finalize"} to get final transcription
    3. Receive transcription results as JSON
    """
    await websocket.accept()

    try:
        # Verify API key if provided
        if ENABLE_AUTH and api_key:
            # Simple validation (in production, verify against database)
            if api_key not in API_KEYS:
                await websocket.send_json({"error": "Invalid API key"})
                await websocket.close()
                return

        audio_chunks = []
        logger.info("WebSocket connection established for streaming transcription")

        while True:
            try:
                # Receive message (can be binary or text)
                data = await websocket.receive()

                if "bytes" in data:
                    # Audio chunk received
                    audio_chunks.append(data["bytes"])
                    await websocket.send_json({
                        "status": "receiving",
                        "chunks_received": len(audio_chunks),
                        "total_size": sum(len(chunk) for chunk in audio_chunks)
                    })

                elif "text" in data:
                    # Command received
                    command = json.loads(data["text"])

                    if command.get("action") == "finalize":
                        # Combine all chunks and transcribe
                        if not audio_chunks:
                            await websocket.send_json({"error": "No audio data received"})
                            continue

                        full_audio = b"".join(audio_chunks)

                        # Save to temp file
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
                            temp_file.write(full_audio)
                            temp_file_path = temp_file.name

                        # Transcribe
                        result = transcribe_file(
                            temp_file_path,
                            command.get("language"),
                            command.get("initial_prompt"),
                            command.get("return_confidence", True)
                        )

                        # Clean up
                        os.unlink(temp_file_path)

                        # Send result
                        await websocket.send_json({
                            "status": "completed",
                            "result": result
                        })

                        # Clear chunks for next session
                        audio_chunks = []

                    elif command.get("action") == "reset":
                        audio_chunks = []
                        await websocket.send_json({"status": "reset", "message": "Audio buffer cleared"})

            except WebSocketDisconnect:
                logger.info("WebSocket disconnected")
                break
            except Exception as e:
                logger.error(f"WebSocket error: {e}")
                await websocket.send_json({"error": str(e)})

    except Exception as e:
        logger.error(f"WebSocket connection error: {e}")
    finally:
        try:
            await websocket.close()
        except:
            pass


# Batch processing endpoint
@app.post(
    "/transcribe/batch",
    dependencies=[Depends(rate_limiter)],
    response_model=None
)
async def transcribe_batch(
    request: Request,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db_session),
    files: List[UploadFile] = File(..., description="Multiple audio files"),
    language: Optional[str] = Form(None),
    webhook_url: Optional[str] = Form(None),
    project_id: Optional[str] = Form(None),
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    """Batch process multiple audio files"""
    correlation_id = request.state.correlation_id

    if len(files) == 0:
        raise HTTPException(status_code=400, detail="No files provided")

    if len(files) > 10:
        raise HTTPException(status_code=400, detail="Maximum 10 files per batch request")

    # Validate language code
    if language and language not in SUPPORTED_LANGUAGES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported language code: {language}. Use /languages endpoint to see supported languages."
        )

    # Validate webhook URL
    if webhook_url and not validate_webhook_url(webhook_url):
        raise HTTPException(
            status_code=400,
            detail="Invalid webhook URL. Must be a valid HTTP or HTTPS URL."
        )

    results = []

    for file in files:
        try:
            content = await file.read()
            file_size = len(content)

            if file_size > MAX_FILE_SIZE:
                results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": f"File too large. Maximum size is {MAX_FILE_SIZE / (1024*1024):.0f}MB"
                })
                continue

            # Create transcription record
            transcription_id = str(uuid.uuid4())
            file_hash = calculate_file_hash(content)

            transcription = Transcription(
                id=transcription_id,
                project_id=project_id,
                correlation_id=f"{correlation_id}:{file.filename}",
                filename=file.filename,
                file_size=file_size,
                file_hash=file_hash,
                status=TranscriptionStatus.PROCESSING,
                language=language,
                webhook_url=webhook_url,
                model_used=MODEL_SIZE,
                created_at=datetime.utcnow(),
                started_at=datetime.utcnow()
            )
            db.add(transcription)
            await db.commit()

            # Queue for async processing if available
            if celery_app:
                task = transcribe_async_task.delay(
                    content,
                    file.filename,
                    transcription_id,
                    language,
                    None,  # initial_prompt
                    False,  # enable_diarization
                    False,  # enable_translation
                    None,  # target_language
                    False,  # enable_sentiment
                    True,  # return_confidence
                    webhook_url,
                    f"{correlation_id}:{file.filename}"
                )
                results.append({
                    "filename": file.filename,
                    "task_id": task.id,
                    "transcription_id": transcription_id,
                    "status": "queued"
                })
            else:
                # Immediate processing
                filename = file.filename or "audio.wav"
                suffix = os.path.splitext(filename)[1] or ".wav"
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                    temp_file.write(content)
                    temp_file_path = temp_file.name

                result = transcribe_file(temp_file_path, language)
                os.unlink(temp_file_path)

                # Update transcription record
                transcription.status = TranscriptionStatus.COMPLETED
                transcription.completed_at = datetime.utcnow()
                transcription.text = result["text"]
                transcription.detected_language = result["language"]
                transcription.segments = result["segments"]
                transcription.processing_time = result["duration"]
                await db.commit()

                results.append({
                    "filename": file.filename,
                    "transcription_id": transcription_id,
                    "status": "completed",
                    "result": result
                })

        except Exception as e:
            results.append({
                "filename": file.filename,
                "status": "error",
                "error": str(e)
            })

    return JSONResponse(content={
        "success": True,
        "correlation_id": correlation_id,
        "total_files": len(files),
        "results": results
    })


# Search transcriptions
@app.get("/transcriptions/search")
async def search_transcriptions(
    query: str,
    project_id: Optional[str] = None,
    limit: int = 10,
    offset: int = 0,
    db: Session = Depends(get_db_session),
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    """
    **Full-Text Search Transcriptions**

    Search across all transcriptions using PostgreSQL full-text search.
    """
    try:
        # Build search query
        search_query = select(Transcription).where(
            Transcription.status == TranscriptionStatus.COMPLETED
        )

        if project_id:
            search_query = search_query.where(Transcription.project_id == project_id)

        # Full-text search on text field
        search_query = search_query.where(
            or_(
                Transcription.text.ilike(f"%{query}%"),
                Transcription.filename.ilike(f"%{query}%")
            )
        ).order_by(Transcription.created_at.desc()).limit(limit).offset(offset)

        result = await db.execute(search_query)
        transcriptions = result.scalars().all()

        return {
            "query": query,
            "total": len(transcriptions),
            "results": [
                {
                    "id": t.id,
                    "filename": t.filename,
                    "text": t.text[:200] + "..." if len(t.text) > 200 else t.text,
                    "language": t.detected_language,
                    "created_at": t.created_at.isoformat(),
                    "duration": t.audio_duration
                }
                for t in transcriptions
            ]
        }

    except Exception as e:
        logger.error(f"Search error: {e}")
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")


# Get transcription by ID
@app.get("/transcriptions/{transcription_id}")
async def get_transcription(
    transcription_id: str,
    db: Session = Depends(get_db_session),
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    """Get transcription by ID"""
    result = await db.execute(
        select(Transcription).where(Transcription.id == transcription_id)
    )
    transcription = result.scalar_one_or_none()

    if not transcription:
        raise HTTPException(status_code=404, detail="Transcription not found")

    return {
        "id": transcription.id,
        "filename": transcription.filename,
        "status": transcription.status,
        "text": transcription.text,
        "language": transcription.detected_language,
        "segments": transcription.segments,
        "confidence_scores": transcription.confidence_scores,
        "diarization": transcription.speakers,
        "translation": transcription.translation,
        "sentiment": transcription.sentiment,
        "processing_time": transcription.processing_time,
        "audio_duration": transcription.audio_duration,
        "created_at": transcription.created_at.isoformat() if transcription.created_at else None,
        "completed_at": transcription.completed_at.isoformat() if transcription.completed_at else None
    }


# Organization management
@app.post("/organizations")
async def create_organization(
    name: str,
    email: str,
    subscription_tier: str = "free",
    db: Session = Depends(get_db_session)
):
    """Create new organization"""
    org_id = str(uuid.uuid4())

    # Determine quota based on tier
    quotas = {
        "free": 1000,
        "starter": 10000,
        "professional": 50000,
        "enterprise": 1000000
    }

    org = Organization(
        id=org_id,
        name=name,
        email=email,
        subscription_tier=SubscriptionTier(subscription_tier),
        monthly_quota=quotas.get(subscription_tier, 1000),
        quota_reset_date=datetime.utcnow() + timedelta(days=30),
        created_at=datetime.utcnow()
    )

    db.add(org)
    await db.commit()

    return {
        "id": org_id,
        "name": name,
        "subscription_tier": subscription_tier,
        "monthly_quota": org.monthly_quota
    }


# Project management
@app.post("/projects")
async def create_project(
    name: str,
    description: Optional[str] = None,
    db: Session = Depends(get_db_session),
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    """Create new project"""
    if not api_key:
        raise HTTPException(status_code=401, detail="API key required")

    project_id = str(uuid.uuid4())

    project = Project(
        id=project_id,
        organization_id=api_key.organization_id,
        name=name,
        description=description,
        created_at=datetime.utcnow()
    )

    db.add(project)
    await db.commit()

    return {
        "id": project_id,
        "name": name,
        "organization_id": api_key.organization_id
    }


# Usage analytics
@app.get("/usage/analytics")
async def get_usage_analytics(
    start_date: Optional[str] = None,
    end_date: Optional[str] = None,
    db: Session = Depends(get_db_session),
    api_key: Optional[APIKey] = Depends(verify_api_key)
):
    """Get usage analytics for organization"""
    if not api_key:
        raise HTTPException(status_code=401, detail="API key required")

    # Build query
    query = select(UsageLog).where(UsageLog.organization_id == api_key.organization_id)

    if start_date:
        query = query.where(UsageLog.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.where(UsageLog.created_at <= datetime.fromisoformat(end_date))

    result = await db.execute(query)
    logs = result.scalars().all()

    # Calculate statistics
    total_duration = sum(log.audio_duration for log in logs)
    total_cost = sum(log.cost for log in logs)
    total_requests = len(logs)

    return {
        "organization_id": api_key.organization_id,
        "period": {
            "start": start_date,
            "end": end_date
        },
        "total_requests": total_requests,
        "total_duration_minutes": total_duration,
        "total_cost_cents": total_cost,
        "features_usage": {
            "diarization": sum(1 for log in logs if log.used_diarization),
            "translation": sum(1 for log in logs if log.used_translation),
            "sentiment": sum(1 for log in logs if log.used_sentiment)
        }
    }


@app.get("/task/{task_id}")
async def get_task_status(task_id: str, api_key: Optional[APIKey] = Depends(verify_api_key)):
    """Get status of async transcription task"""
    if not celery_app:
        raise HTTPException(
            status_code=503,
            detail="Async processing not available. Celery is not configured."
        )

    if not task_id or not task_id.strip():
        raise HTTPException(
            status_code=400,
            detail="Invalid task_id provided"
        )

    try:
        from celery.result import AsyncResult
        task = AsyncResult(task_id, app=celery_app)

        if task.ready():
            if task.successful():
                return {
                    "task_id": task_id,
                    "status": "completed",
                    "result": task.result
                }
            else:
                return {
                    "task_id": task_id,
                    "status": "failed",
                    "error": str(task.info)
                }
        else:
            return {
                "task_id": task_id,
                "status": "processing"
            }
    except Exception as e:
        logger.error(f"Error retrieving task status for {task_id}: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error retrieving task status: {str(e)}"
        )


@app.get("/languages")
async def get_supported_languages():
    """Get list of supported language codes"""
    return {
        "supported_languages": sorted(list(SUPPORTED_LANGUAGES)),
        "count": len(SUPPORTED_LANGUAGES),
        "note": "Language codes follow ISO 639-1 standard. Leave language parameter empty for auto-detection."
    }


@app.get("/health")
async def health_check():
    """Health check endpoint with system status"""
    redis_status = "connected" if redis_client else "disconnected"
    celery_status = "available" if celery_app else "unavailable"
    db_status = "connected" if check_db_connection() else "disconnected"

    return {
        "status": "healthy",
        "model": MODEL_SIZE,
        "device": DEVICE,
        "service": "Enterprise STT API v3.0",
        "version": "3.0.0",
        "max_file_size_mb": MAX_FILE_SIZE / (1024*1024),
        "supported_languages_count": len(SUPPORTED_LANGUAGES),
        "features": {
            "authentication": ENABLE_AUTH,
            "rate_limiting": True,
            "caching": redis_status,
            "async_processing": celery_status,
            "database": db_status,
            "metrics": PROMETHEUS_AVAILABLE,
            "gpu_support": USE_GPU,
            "speaker_diarization": DIARIZATION_AVAILABLE,
            "cloud_storage": {
                "s3": S3_AVAILABLE,
                "gcs": GCS_AVAILABLE,
                "azure": AZURE_AVAILABLE
            },
            "export_formats": {
                "docx": DOCX_AVAILABLE,
                "pdf": PDF_AVAILABLE
            },
            "translation": TRANSFORMERS_AVAILABLE,
            "sentiment_analysis": sentiment_analyzer is not None,
            "real_time_streaming": True
        }
    }


@app.get("/metrics")
async def metrics():
    """Prometheus metrics endpoint"""
    if not PROMETHEUS_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="Metrics not available. Prometheus client is not installed."
        )

    from fastapi import Response
    return Response(content=generate_latest(), media_type=CONTENT_TYPE_LATEST)


@app.get("/analytics")
async def get_analytics(api_key: Optional[APIKey] = Depends(verify_api_key)):
    """Get usage analytics"""
    avg_duration = (
        usage_stats["total_duration"] / usage_stats["successful_requests"]
        if usage_stats["successful_requests"] > 0
        else 0
    )

    cache_hit_rate = (
        usage_stats["cache_hits"] / (usage_stats["cache_hits"] + usage_stats["cache_misses"])
        if (usage_stats["cache_hits"] + usage_stats["cache_misses"]) > 0
        else 0
    )

    return {
        "total_requests": usage_stats["total_requests"],
        "successful_requests": usage_stats["successful_requests"],
        "failed_requests": usage_stats["failed_requests"],
        "success_rate": (
            usage_stats["successful_requests"] / usage_stats["total_requests"]
            if usage_stats["total_requests"] > 0
            else 0
        ),
        "average_processing_time": avg_duration,
        "cache_statistics": {
            "hits": usage_stats["cache_hits"],
            "misses": usage_stats["cache_misses"],
            "hit_rate": cache_hit_rate
        }
    }


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard():
    """Serve the web dashboard"""
    try:
        with open("static/index.html", "r") as f:
            return f.read()
    except FileNotFoundError:
        return HTMLResponse(content="<h1>Dashboard not found</h1><p>Please ensure static/index.html exists</p>", status_code=404)


@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "service": "Enterprise Speech-to-Text API v3.0",
        "version": "3.0.0",
        "model": MODEL_SIZE,
        "device": DEVICE,
        "max_file_size_mb": MAX_FILE_SIZE / (1024*1024),
        "features": {
            "multi_tenant": True,
            "speaker_diarization": DIARIZATION_AVAILABLE,
            "auto_language_detection": True,
            "confidence_scores": True,
            "custom_prompts": True,
            "translation": TRANSFORMERS_AVAILABLE,
            "sentiment_analysis": sentiment_analyzer is not None,
            "real_time_streaming": True,
            "cloud_storage": {
                "s3": S3_AVAILABLE,
                "gcs": GCS_AVAILABLE,
                "azure": AZURE_AVAILABLE
            },
            "export_formats": ["json", "srt", "vtt", "txt", "docx", "pdf"],
            "batch_processing": True,
            "webhooks_with_retry": True,
            "database_storage": True,
            "full_text_search": True,
            "usage_quotas": True,
            "caching": redis_client is not None,
            "async_queue": celery_app is not None,
            "gpu_support": USE_GPU,
            "metrics": PROMETHEUS_AVAILABLE
        },
        "endpoints": {
            "dashboard": "/dashboard (GET) - Web Dashboard UI",
            "transcribe": "/transcribe (POST) - Advanced transcription with all features",
            "transcribe_batch": "/transcribe/batch (POST) - Batch process multiple files",
            "transcribe_stream": "/ws/transcribe (WebSocket) - Real-time streaming",
            "search": "/transcriptions/search (GET) - Full-text search",
            "get_transcription": "/transcriptions/{id} (GET) - Get transcription by ID",
            "task_status": "/task/{task_id} (GET) - Get async task status",
            "create_org": "/organizations (POST) - Create organization",
            "create_project": "/projects (POST) - Create project",
            "usage_analytics": "/usage/analytics (GET) - Usage analytics",
            "languages": "/languages (GET) - Get supported language codes",
            "health": "/health (GET) - Health check",
            "metrics": "/metrics (GET) - Prometheus metrics",
            "analytics": "/analytics (GET) - System analytics",
            "docs": "/docs - Interactive API documentation",
            "redoc": "/redoc - Alternative API documentation"
        },
        "documentation": "/docs"
    }
